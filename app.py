import streamlit as st
import google.generativeai as genai
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import PyPDF2
import pandas as pd
from datetime import datetime
import time
import random
from io import BytesIO
from docx import Document

# --- 1. Ø¥Ø¹Ø¯Ø§Ø¯ Ø§Ù„ØµÙØ­Ø© (Ù…Ù‡Ù…: Ø§Ù„Ø­Ø§Ù„Ø© collapsed) ---
st.set_page_config(
    page_title="Virtual Supervisor v2", 
    layout="wide", 
    page_icon="ğŸ“",
    initial_sidebar_state="collapsed"
)

# --- 2. Ù…Ù†Ø·Ù‚ Ø²Ø± Ø§Ù„Ù‚Ø§Ø¦Ù…Ø© Ù„Ù„Ù‡Ø§ØªÙ (Session State) ---
if 'mobile_menu_open' not in st.session_state:
    st.session_state.mobile_menu_open = False

def open_menu():
    st.session_state.mobile_menu_open = True

def close_menu():
    st.session_state.mobile_menu_open = False

# ==========================================
# ğŸ¨ CSS: Ø§Ù„ØªØµÙ…ÙŠÙ… + Ø¥Ø¬Ø¨Ø§Ø± Ø§Ù„Ù‚Ø§Ø¦Ù…Ø© Ø¹Ù„Ù‰ Ø§Ù„Ø¸Ù‡ÙˆØ±
# ==========================================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@400;500;700;900&family=Poppins:wght@300;400;600;900&display=swap');
    
    html, body, [class*="css"] { font-family: 'Poppins', 'Tajawal', sans-serif; }

    [data-testid="stAppViewContainer"] {
        background: linear-gradient(135deg, #fdfbfb 0%, #e3f2fd 100%);
        background-attachment: fixed;
    }
    
    /* Ø¥Ø®ÙØ§Ø¡ Ø§Ù„Ù‡ÙŠØ¯Ø± ÙˆØ´Ø±ÙŠØ· Ø§Ù„Ø£Ø¯ÙˆØ§Øª */
    [data-testid="stHeader"], [data-testid="stToolbar"] {
        display: none !important;
        visibility: hidden !important;
        height: 0px !important;
    }
    
    /* ØªÙ†Ø³ÙŠÙ‚ Ø²Ø± ÙØªØ­ Ø§Ù„Ù‚Ø§Ø¦Ù…Ø© (ÙŠØ¸Ù‡Ø± ÙÙŠ Ø§Ù„Ù…ÙˆØ¨Ø§ÙŠÙ„) */
    div.stButton > button.open-menu-btn {
        background-color: #1565c0;
        color: white;
        border-radius: 10px;
        padding: 5px 15px;
        border: none;
        box-shadow: 0 2px 5px rgba(0,0,0,0.2);
    }

    /* --- IMPORTANT: Mobile Sidebar Logic --- */
    /* Ø¨Ø´ÙƒÙ„ Ø§ÙØªØ±Ø§Ø¶ÙŠ Ù†Ø®ÙÙŠ Ø²Ø± Ø§Ù„Ø¥ØºÙ„Ø§Ù‚ Ø§Ù„Ø£ØµÙ„ÙŠ */
    [data-testid="stSidebarCollapseButton"] { display: none !important; }

    /* ØªØ­Ø³ÙŠÙ† Ø´ÙƒÙ„ Ø§Ù„Ù‚Ø§Ø¦Ù…Ø© */
    [data-testid="stSidebar"] {
        background-color: #ffffff !important;
        border-right: 1px solid #e0e0e0;
        padding-top: 20px !important;
    }

    .block-container { padding-top: 1rem !important; }

    /* ÙƒÙ„Ø§Ø³Ø§Øª Ø§Ù„ØªØµÙ…ÙŠÙ… Ø§Ù„Ù‚Ø¯ÙŠÙ…Ø© */
    .hero-box { text-align: center; padding: 60px 20px; background: linear-gradient(135deg, #e3f2fd 0%, #bbdefb 100%); border-radius: 25px; margin-bottom: 40px; border: 1px solid #90caf9; box-shadow: 0 10px 30px rgba(33, 150, 243, 0.15); }
    .hero-title { font-size: 3.5rem; font-weight: 900; color: #1565c0; margin-bottom: 5px; letter-spacing: -1px; }
    .hero-slogan { font-family: 'Poppins', sans-serif; font-size: 1.4rem; font-weight: 700; color: #1976d2; text-transform: uppercase; letter-spacing: 2px; margin-top: 10px; }
    .global-header { text-align: center; padding-bottom: 20px; margin-bottom: 30px; border-bottom: 2px solid rgba(0,0,0,0.05); }
    .main-title { font-family: 'Poppins', sans-serif; font-size: 3rem; font-weight: 900; color: #1565c0; margin: 0; letter-spacing: -1px; line-height: 1.1; }
    .fixed-slogan { font-family: 'Poppins', sans-serif; background: -webkit-linear-gradient(45deg, #1e3c72, #2a5298); -webkit-background-clip: text; -webkit-text-fill-color: transparent; font-size: 1.6rem; font-weight: 800; text-transform: uppercase; letter-spacing: 3px; margin-top: 5px; }
    .info-section { background: white; padding: 30px; border-radius: 20px; margin-bottom: 30px; border-left: 5px solid #2196f3; box-shadow: 0 5px 15px rgba(0,0,0,0.05); }
    .service-card { background: white; padding: 25px; border-radius: 15px; text-align: center; box-shadow: 0 5px 15px rgba(0,0,0,0.05); border: 1px solid #e3f2fd; height: 100%; transition: transform 0.3s; }
    .contact-section { background: #f1f8ff; padding: 30px; border-radius: 20px; margin-top: 40px; border: 1px solid #d1e9ff; }
    .plan-card { background: white; border-radius: 15px; padding: 20px; text-align: center; border: 1px solid #eee; box-shadow: 0 5px 15px rgba(0,0,0,0.05); height: 100%; display: flex; flex-direction: column; justify-content: space-between; }
    .price-tag { font-size: 2rem; font-weight: 900; color: #2c3e50; margin: 15px 0; }
    .result-card { background: white; padding: 30px; border-radius: 20px; margin-bottom: 20px; box-shadow: 0 5px 15px rgba(0,0,0,0.05); }
    .integrity-box { background: #fff3cd; color: #856404; border: 1px solid #ffeeba; padding: 15px; border-radius: 12px; margin-bottom: 25px; display: flex; align-items: center; gap: 15px; }
    .stButton button { border-radius: 50px; font-weight: bold; background: linear-gradient(90deg, #00d2ff 0%, #3a7bd5 100%); color: white; border: none; }
    [data-testid="stChatMessage"] { background: white; border-radius: 15px; }
</style>
""", unsafe_allow_html=True)

# --- CSS Injection Based on State ---
# Ù‡Ø°Ø§ Ø§Ù„ÙƒÙˆØ¯ Ù‡Ùˆ Ø§Ù„Ù…Ø³Ø¤ÙˆÙ„ Ø¹Ù† Ø¥Ø¬Ø¨Ø§Ø± Ø§Ù„Ù‚Ø§Ø¦Ù…Ø© Ø¹Ù„Ù‰ Ø§Ù„Ø¸Ù‡ÙˆØ± ÙÙŠ Ø§Ù„Ù…ÙˆØ¨Ø§ÙŠÙ„
if st.session_state.mobile_menu_open:
    st.markdown("""
        <style>
        @media (max-width: 991px) {
            [data-testid="stSidebar"] {
                display: block !important;
                width: 100% !important;
                z-index: 999999;
                position: fixed;
                height: 100vh;
                top: 0;
                left: 0;
            }
        }
        </style>
    """, unsafe_allow_html=True)
else:
    # Ø¥Ø®ÙØ§Ø¡ Ø§Ù„Ù‚Ø§Ø¦Ù…Ø© ÙÙŠ Ø§Ù„Ù…ÙˆØ¨Ø§ÙŠÙ„ Ø¥Ø°Ø§ Ù„Ù… ÙŠØªÙ… ØªÙØ¹ÙŠÙ„ Ø§Ù„Ø²Ø±
    st.markdown("""
        <style>
        @media (max-width: 991px) {
            [data-testid="stSidebar"] {
                display: none;
            }
        }
        </style>
    """, unsafe_allow_html=True)

# --- Ø²Ø± ÙØªØ­ Ø§Ù„Ù‚Ø§Ø¦Ù…Ø© (ÙŠØ¸Ù‡Ø± ÙÙ‚Ø· Ø¥Ø°Ø§ ÙƒØ§Ù†Øª Ù…ØºÙ„Ù‚Ø©) ---
if not st.session_state.mobile_menu_open:
    col_menu, col_space = st.columns([1, 10])
    with col_menu:
        if st.button("â˜°", key="main_open_btn", help="Open Menu"):
            open_menu()
            st.rerun()

# ==========================================
# ğŸŒ UI Dictionary & Config
# ==========================================
UI_TEXT = {
    "English": {
        "dir": "ltr", "align": "left",
        "chat_title": "Chat with Supervisor",
        "sidebar_settings": "âš™ï¸ Settings",
        "field_label": "Academic Field",
        "level_label": "Academic Level",
        "task_label": "Select Service",
        "history_label": "ğŸ“‚ Research History",
        "input_ph": "Enter your research topic or text here...",
        "ref_ph": "Paste your references list here...",
        "format_label": "Select Citation Style",
        "citation_styles": ["APA 7", "MLA 9", "Chicago", "Harvard", "IEEE"],
        "file_ph": "Upload PDF Document",
        "exec_btn": "âœ¨ Generate Magic",
        "save_btn": "ğŸ’¾ Save to History",
        "dl_btn": "ğŸ“¥ Download (Word Doc)",
        "warn_title": "IMPORTANT ACADEMIC INTEGRITY NOTICE",
        "warn_msg": "This tool is an AI assistant designed to guide and structure your thoughts, NOT to write your thesis for you. Copying content directly is considered plagiarism.",
        "upgrade_btn": "ğŸ”“ Upgrade to Unlock Full Plan",
        "pay_title": "âœ¨ Upgrade to Premium",
        "pay_pitch_title": "Why Subscribe?",
        "pay_pitch_body": "Specially tuned for academic research standards. Get deep analysis, APA citations, and structured plans.",
        "plans": {"1": "Monthly", "6": "6 Months", "12": "Yearly"},
        "plan_desc": {"1": "Flexible start", "6": "Best Value!", "12": "Full commitment"},
        "pay_msg": "ğŸ”’ Preview Mode. Upgrade to see full content.",
        "select_btn": "Select",
        "pay_success": "Payment Sent! Confirmation email coming soon.",
        "pay_error": "Please enter transaction ID.",
        "cancel_btn": "ğŸ”™ Return to Workspace",
        "fields": ["Science & Tech", "Medical", "Law", "Economics", "Arts", "Humanities", "Islamic", "Architecture"],
        "levels": ["Master's", "PhD", "Researcher"],
        "tasks": {
            "Discuss Research Topic (Free)": "discuss_topic",
            "Research Plan Proposal": "structure",
            "Suggest Academic References": "references",
            "Format Bibliography (APA/MLA)": "formatting",
            "Scientific Proofreading": "proofread",
            "Analyze & Summarize Reference": "analyze"
        }
    },
    "FranÃ§ais": {
        "dir": "ltr", "align": "left",
        "chat_title": "Discuter avec Superviseur",
        "sidebar_settings": "ParamÃ¨tres",
        "field_label": "Domaine",
        "level_label": "Niveau",
        "task_label": "Service",
        "history_label": "ğŸ“‚ Historique",
        "input_ph": "Saisissez votre sujet ici...",
        "ref_ph": "Collez votre liste de rÃ©fÃ©rences ici...",
        "format_label": "Style de citation",
        "citation_styles": ["APA 7", "MLA 9", "Chicago", "Harvard", "IEEE"],
        "file_ph": "TÃ©lÃ©charger PDF",
        "exec_btn": "âœ¨ Lancer l'Analyse",
        "save_btn": "ğŸ’¾ Sauvegarder",
        "dl_btn": "ğŸ“¥ TÃ©lÃ©charger (Word)",
        "warn_title": "AVIS D'INTÃ‰GRITÃ‰ ACADÃ‰MIQUE",
        "warn_msg": "Cet outil est un assistant conÃ§u pour vous guider, PAS pour rÃ©diger Ã  votre place. Le copier-coller direct est considÃ©rÃ© comme du plagiat.",
        "upgrade_btn": "ğŸ”“ Passer en Premium",
        "pay_title": "âœ¨ Passer en Premium",
        "pay_pitch_title": "Pourquoi s'abonner ?",
        "pay_pitch_body": "SpÃ©cialisÃ© pour les normes acadÃ©miques. Obtenez des analyses profondes et des plans structurÃ©s.",
        "plans": {"1": "Mensuel", "6": "6 Mois", "12": "Annuel"},
        "plan_desc": {"1": "Flexible", "6": "Meilleure Valeur", "12": "Annuel"},
        "pay_msg": "ğŸ”’ Mode AperÃ§u. Abonnez-vous pour tout voir.",
        "select_btn": "Choisir",
        "pay_success": "Paiement envoyÃ© ! Confirmation par e-mail bientÃ´t.",
        "pay_error": "Entrez le numÃ©ro.",
        "cancel_btn": "ğŸ”™ Retour",
        "fields": ["Sciences & Tech", "MÃ©dical", "Droit", "Ã‰conomie", "Lettres", "Humaines", "Islamiques", "Architecture"],
        "levels": ["Master", "Doctorat", "Chercheur"],
        "tasks": {
            "Discuter du Sujet (Gratuit)": "discuss_topic",
            "Proposition de Plan": "structure",
            "Suggestion de RÃ©fÃ©rences": "references",
            "Mise en forme Bibliographie": "formatting",
            "Correction AcadÃ©mique": "proofread",
            "Analyse et RÃ©sumÃ© de RÃ©fÃ©rence": "analyze"
        }
    },
    "Ø§Ù„Ø¹Ø±Ø¨ÙŠØ©": {
        "dir": "rtl", "align": "right",
        "chat_title": "ØªØ­Ø¯Ø« Ù…Ø¹ Ù…Ø´Ø±ÙÙƒ",
        "sidebar_settings": "Ø¥Ø¹Ø¯Ø§Ø¯Ø§Øª Ø§Ù„Ø¨Ø­Ø«",
        "field_label": "Ø§Ù„Ù…Ø¬Ø§Ù„ Ø§Ù„Ø¯Ø±Ø§Ø³ÙŠ",
        "level_label": "Ø§Ù„Ù…Ø³ØªÙˆÙ‰ Ø§Ù„Ø£ÙƒØ§Ø¯ÙŠÙ…ÙŠ",
        "task_label": "Ø§Ù„Ø®Ø¯Ù…Ø© Ø§Ù„Ù…Ø·Ù„ÙˆØ¨Ø©",
        "history_label": "ğŸ“‚ Ø£Ø±Ø´ÙŠÙ Ø£Ø¨Ø­Ø§Ø«ÙŠ",
        "input_ph": "Ø§ÙƒØªØ¨ Ù…ÙˆØ¶ÙˆØ¹ Ø§Ù„Ø¨Ø­Ø« Ø£Ùˆ Ø§Ù„Ù†Øµ Ù‡Ù†Ø§...",
        "ref_ph": "Ø£Ù„ØµÙ‚ Ù‚Ø§Ø¦Ù…Ø© Ø§Ù„Ù…Ø±Ø§Ø¬Ø¹ Ù‡Ù†Ø§...",
        "format_label": "Ù†Ø¸Ø§Ù… Ø§Ù„ØªÙˆØ«ÙŠÙ‚ Ø§Ù„Ù…Ø·Ù„ÙˆØ¨",
        "citation_styles": ["APA 7", "MLA 9", "Chicago", "Harvard", "IEEE"],
        "file_ph": "Ø±ÙØ¹ Ù…Ù„Ù Ø§Ù„Ù…Ø±Ø¬Ø¹ (PDF)",
        "exec_btn": "âœ¨ Ø§Ø¨Ø¯Ø£ Ø§Ù„Ù…Ø¹Ø§Ù„Ø¬Ø©",
        "save_btn": "ğŸ’¾ Ø­ÙØ¸ ÙÙŠ Ø§Ù„Ø£Ø±Ø´ÙŠÙ",
        "dl_btn": "ğŸ“¥ ØªØ­Ù…ÙŠÙ„ Ù…Ù„Ù ÙˆÙˆØ±Ø¯",
        "warn_title": "ØªÙ†Ø¨ÙŠÙ‡ Ù‡Ø§Ù… Ø­ÙˆÙ„ Ø§Ù„Ø£Ù…Ø§Ù†Ø© Ø§Ù„Ø¹Ù„Ù…ÙŠØ©",
        "warn_msg": "ØªÙ… ØªØµÙ…ÙŠÙ… Ù‡Ø°Ø§ Ø§Ù„Ù…Ø´Ø±Ù Ø§Ù„Ø§ÙØªØ±Ø§Ø¶ÙŠ Ù„ÙŠÙƒÙˆÙ† Ù…ÙˆØ¬Ù‡Ø§Ù‹ ÙˆÙ…Ø³Ø§Ø¹Ø¯Ø§Ù‹ Ù„Ùƒ Ù„ØªÙ†Ø¸ÙŠÙ… Ø£ÙÙƒØ§Ø±ÙƒØŒ ÙˆÙ„ÙŠØ³ Ù„ÙŠÙ‚ÙˆÙ… Ø¨ÙƒØªØ§Ø¨Ø© Ø§Ù„Ø¨Ø­Ø« Ù†ÙŠØ§Ø¨Ø© Ø¹Ù†Ùƒ. Ø§Ù„Ù†Ø³Ø® ÙˆØ§Ù„Ù„ØµÙ‚ Ø§Ù„Ù…Ø¨Ø§Ø´Ø± ÙŠØ¹ØªØ¨Ø± Ø³Ø±Ù‚Ø© Ø¹Ù„Ù…ÙŠØ©.",
        "upgrade_btn": "ğŸ”“ Ø§Ø´ØªØ±Ùƒ Ø§Ù„Ø¢Ù† Ù„Ø¥Ø¸Ù‡Ø§Ø± Ø§Ù„Ø®Ø·Ø© ÙƒØ§Ù…Ù„Ø©",
        "pay_msg": "ğŸ”’ Ø§Ø´ØªØ±Ùƒ Ø§Ù„Ø¢Ù† Ù„Ø¥ÙƒÙ…Ø§Ù„ Ø§Ù„Ù‚Ø±Ø§Ø¡Ø©",
        "pay_title": "âœ¨ ØªØ±Ù‚ÙŠØ© Ø§Ù„Ø¹Ø¶ÙˆÙŠØ© (Premium)",
        "pay_pitch_title": "Ù„Ù…Ø§Ø°Ø§ ØªØ´ØªØ±ÙƒØŸ",
        "pay_pitch_body": "ØªÙ… ØªØ¯Ø±ÙŠØ¨ Ø§Ù„Ù…Ø´Ø±Ù Ø§Ù„Ø§ÙØªØ±Ø§Ø¶ÙŠ Ø®ØµÙŠØµØ§Ù‹ Ù„Ù„Ù…Ø¹Ø§ÙŠÙŠØ± Ø§Ù„Ø£ÙƒØ§Ø¯ÙŠÙ…ÙŠØ©. Ø§Ø­ØµÙ„ Ø¹Ù„Ù‰ Ø®Ø·Ø· ÙƒØ§Ù…Ù„Ø©ØŒ ØªØ­Ù„ÙŠÙ„ Ø¹Ù…ÙŠÙ‚ØŒ ÙˆÙ…Ø±Ø§ÙÙ‚Ø© Ø¯Ø§Ø¦Ù…Ø©.",
        "plans": {"1": "Ø´Ù‡Ø±ÙŠ", "6": "6 Ø£Ø´Ù‡Ø±", "12": "Ø³Ù†ÙˆÙŠ"},
        "plan_desc": {"1": "Ø¨Ø¯Ø§ÙŠØ© Ù…Ø±Ù†Ø©", "6": "Ø§Ù„Ø£ÙƒØ«Ø± Ø·Ù„Ø¨Ø§Ù‹!", "12": "Ø§Ù„ØªØ²Ø§Ù… Ø³Ù†ÙˆÙŠ"},
        "select_btn": "Ø§Ø®ØªØ±",
        "pay_success": "ØªÙ… Ø§Ø±Ø³Ø§Ù„ Ø§Ù„Ø§Ø´ØªØ±Ø§ÙƒØŒ Ø³ÙŠØµÙ„Ùƒ Ø¨Ø±ÙŠØ¯ Ø¥Ù„ÙƒØªØ±ÙˆÙ†ÙŠ Ù„ØªØ£ÙƒÙŠØ¯ ØªÙØ¹ÙŠÙ„ Ø­Ø³Ø§Ø¨Ùƒ.",
        "pay_error": "Ø§Ù„Ø±Ø¬Ø§Ø¡ Ø¥Ø¯Ø®Ø§Ù„ Ø±Ù‚Ù… Ø§Ù„ÙˆØµÙ„.",
        "cancel_btn": "ğŸ”™ Ø§Ù„Ø¹ÙˆØ¯Ø© Ù„Ù…Ø³Ø§Ø­Ø© Ø§Ù„Ø¹Ù…Ù„",
        "fields": ["Ø§Ù„Ø¹Ù„ÙˆÙ… ÙˆØ§Ù„ØªÙƒÙ†ÙˆÙ„ÙˆØ¬ÙŠØ§", "Ø§Ù„Ø·Ø¨ ÙˆØ§Ù„ØµÙŠØ¯Ù„Ø©", "Ø§Ù„Ø­Ù‚ÙˆÙ‚ ÙˆØ§Ù„Ø³ÙŠØ§Ø³Ø©", "Ø§Ù„Ø§Ù‚ØªØµØ§Ø¯", "Ø§Ù„Ø¢Ø¯Ø§Ø¨ ÙˆØ§Ù„Ù„ØºØ§Øª", "Ø§Ù„Ø¹Ù„ÙˆÙ… Ø§Ù„Ø¥Ù†Ø³Ø§Ù†ÙŠØ©", "Ø§Ù„Ø¹Ù„ÙˆÙ… Ø§Ù„Ø¥Ø³Ù„Ø§Ù…ÙŠØ©", "Ø§Ù„Ø¹Ù…Ø±Ø§Ù†"],
        "levels": ["Ù…Ø§Ø³ØªØ±", "Ø¯ÙƒØªÙˆØ±Ø§Ù‡", "Ø¨Ø§Ø­Ø« Ø£ÙƒØ§Ø¯ÙŠÙ…ÙŠ"],
        "tasks": {
            "Ù…Ù†Ø§Ù‚Ø´Ø© Ù…ÙˆØ¶ÙˆØ¹ Ø§Ù„Ø¨Ø­Ø« (Ù…Ø¬Ø§Ù†ÙŠ)": "discuss_topic",
            "Ø§Ù‚ØªØ±Ø§Ø­ Ø®Ø·Ø© Ø¹Ù…Ù„": "structure",
            "Ø§Ù‚ØªØ±Ø§Ø­ Ù…Ø±Ø§Ø¬Ø¹ Ø§ÙƒØ§Ø¯ÙŠÙ…ÙŠØ©": "references",
            "ØªÙ†Ø³ÙŠÙ‚ ÙˆØªÙ†Ø¸ÙŠÙ… Ø§Ù„Ù…Ø±Ø§Ø¬Ø¹ (APA)": "formatting",
            "ØªØ¯Ù‚ÙŠÙ‚ Ø¹Ù„Ù…ÙŠ": "proofread",
            "ØªØ­Ù„ÙŠÙ„ ÙˆØªÙ„Ø®ÙŠØµ Ù…Ø±Ø¬Ø¹": "analyze"
        }
    }
}

# ==========================================
# ğŸ”§ System Setup
# ==========================================
ADMIN_EMAIL = "souad.belkhanousse@gmail.com"

if "init" not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.user_info = {}
    st.session_state.chat_history = []
    st.session_state.current_context = ""
    st.session_state.res_restored = None
    st.session_state.restored_task = ""
    st.session_state.last_res = None
    st.session_state.show_payment_page = False
    st.session_state.selected_plan = None
    st.session_state.page_state = "landing"
    st.session_state.init = True

def rain_graduation_caps():
    js = """<script>function createCap() {const cap = document.createElement('div');cap.innerText = 'ğŸ“';cap.style.position = 'fixed';cap.style.left = Math.random() * 100 + 'vw';cap.style.bottom = '-50px';cap.style.fontSize = (Math.random() * 20 + 30) + 'px';cap.style.animation = 'floatUp ' + (Math.random() * 3 + 2) + 's linear';cap.style.zIndex = '99999999';document.body.appendChild(cap);setTimeout(() => { cap.remove(); }, 5000);}for(let i=0; i<50; i++) { setTimeout(createCap, i * 100); }</script>"""
    st.components.v1.html(js, height=0)

def set_archive(content, task):
    st.session_state.res_restored = content
    st.session_state.restored_task = task
    st.session_state.current_context = f"ARCHIVED: {content}"

def close_archive(): st.session_state.res_restored = None
def logout():
    for k in list(st.session_state.keys()): del st.session_state[k]
    st.rerun()

def go_to_auth(): st.session_state.page_state = "login"; st.rerun()
def go_to_payment(): st.session_state.show_payment_page = True

@st.cache_resource
def get_db():
    scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
    creds = ServiceAccountCredentials.from_json_keyfile_dict(dict(st.secrets["gcp_service_account"]), scope)
    return gspread.authorize(creds)

def get_sheet(name): return get_db().open("dz_supervisor_users").worksheet(name)

def check_login(e, p):
    try:
        sh = get_sheet("users"); rec = sh.get_all_records()
        for u in rec:
            if str(u.get('username','')).strip().lower() == e.strip().lower() and str(u.get('password','')).strip() == p.strip():
                return True, u
        return False, None
    except: return False, None

def register_user(e, p, n):
    try:
        sh = get_sheet("users"); e = e.strip().lower()
        status = "active" if e == ADMIN_EMAIL.strip().lower() else "pending"
        sh.append_row([e, p.strip(), n, status, "2025-12-31"])
        return True, "Success"
    except: return False, "Error"

def submit_payment(email, ref, plan):
    try:
        sh = get_sheet("users"); cell = sh.find(email); sh.update_cell(cell.row, 4, "review"); return True
    except: return False

def save_research(email, task, content):
    try: sh = get_sheet("history"); sh.append_row([email, task, content, datetime.now().strftime("%Y-%m-%d %H:%M")]); return True
    except: return False

def get_history(email):
    try: sh = get_sheet("history"); return [r for r in sh.get_all_records() if str(r.get('email','')).lower() == email.strip().lower()]
    except: return []

def create_word_docx(content, title="Result"):
    doc = Document(); doc.add_heading(title, 0); doc.add_paragraph(content); bio = BytesIO(); doc.save(bio); return bio.getvalue()

# ==========================================
# ğŸ  LANDING PAGE
# ==========================================
if not st.session_state.logged_in and st.session_state.page_state == "landing":
    # --- Hero Section ---
    st.markdown("""
    <div class="hero-box">
        <img src="https://cdn-icons-png.flaticon.com/512/3135/3135768.png" width="120" style="margin-bottom:15px;">
        <h1 class="hero-title">Virtual Supervisor v2</h1>
        <div class="hero-slogan">Research Smarter, Not Harder</div>
    </div>
    """, unsafe_allow_html=True)
    
    # --- Who is VS? ---
    st.markdown("""
    <div class="info-section" style="direction:rtl; text-align:right;">
        <h2 style="color:#1565c0; text-align:center; margin-bottom:20px;">ğŸ“ Ù…Ù† Ù‡Ùˆ Ø§Ù„Ù…Ø´Ø±Ù Ø§Ù„Ø§ÙØªØ±Ø§Ø¶ÙŠØŸ</h2>
        <div style="background:#e3f2fd; padding:20px; border-radius:10px; margin-bottom:20px;">
            <p style="color:#0d47a1; font-weight:bold; text-align:center;">" Ù„Ù† ØªØ¶Ø·Ø± Ø¥Ù„Ù‰ Ø¥ÙŠÙ‚Ø§Ù Ø¨Ø­Ø«Ùƒ ÙÙŠ Ø§Ù†ØªØ¸Ø§Ø± Ø§Ù„Ù…ÙˆØ§Ø¹ÙŠØ¯ Ø¨Ø¹Ø¯ Ø§Ù„ÙŠÙˆÙ…! "</p>
        </div>
        <div class="bilingual-box">
            <div class="info-text-en" style="direction:ltr; text-align:left; margin-bottom:15px; padding-bottom:15px; border-bottom:1px dashed #ddd;">
                <b>Virtual Supervisor</b> is an advanced AI system trained specifically on academic methodologies. It acts as your 24/7 mentor.
            </div>
            <div class="info-text-ar">
                <b>Ø§Ù„Ù…Ø´Ø±Ù Ø§Ù„Ø§ÙØªØ±Ø§Ø¶ÙŠ</b> Ù‡Ùˆ Ù†Ø¸Ø§Ù… Ø°ÙƒÙŠ Ù…ØªØ·ÙˆØ± ØªÙ… ØªØ¯Ø±ÙŠØ¨Ù‡ Ø®ØµÙŠØµØ§Ù‹ Ø¹Ù„Ù‰ Ø§Ù„Ù…Ù†Ù‡Ø¬ÙŠØ§Øª Ø§Ù„Ø£ÙƒØ§Ø¯ÙŠÙ…ÙŠØ©. ÙŠØ¹Ù…Ù„ ÙƒÙ…ÙˆØ¬Ù‡ Ø´Ø®ØµÙŠ Ù…ØªØ§Ø­ 24/7 Ù„Ù…Ø³Ø§Ø¹Ø¯ØªÙƒ ÙÙŠ ØªØ¬Ø§ÙˆØ² Ø¹Ù‚Ø¨Ø§Øª Ø§Ù„ÙƒØªØ§Ø¨Ø© ÙˆØ§Ù„ØªØ­Ù„ÙŠÙ„ Ø§Ù„ÙÙ†ÙŠ ÙÙˆØ±Ø§Ù‹.
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("<h2 style='text-align:center; color:#0d47a1; margin: 40px 0;'>Ø®Ø¯Ù…Ø§ØªÙ†Ø§ Ø§Ù„Ù…ØªÙ…ÙŠØ²Ø©</h2>", unsafe_allow_html=True)
    
    # --- Services Grid ---
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.markdown("""
        <div class="service-card">
            <span class="srv-icon">ğŸ“‹</span>
            <div class="srv-title">Ø§Ù‚ØªØ±Ø§Ø­ Ø§Ù„Ø®Ø·Ø·</div>
            <div class="srv-desc">Ø¨Ù†Ø§Ø¡ Ù‡ÙŠÙƒÙ„ Ø¨Ø­Ø«ÙŠ Ù…ØªÙƒØ§Ù…Ù„ (ÙØµÙˆÙ„ ÙˆÙ…Ø¨Ø§Ø­Ø«) Ø¨Ù…Ù†Ù‡Ø¬ÙŠØ© Ø¹Ù„Ù…ÙŠØ©.</div>
        </div>""", unsafe_allow_html=True)
    with c2:
        st.markdown("""
        <div class="service-card">
            <span class="srv-icon">ğŸ“š</span>
            <div class="srv-title">ØªÙ†Ø³ÙŠÙ‚ Ø§Ù„Ù…Ø±Ø§Ø¬Ø¹</div>
            <div class="srv-desc">Ø§Ù‚ØªØ±Ø§Ø­ ÙˆØ¶Ø¨Ø· Ø§Ù„Ù…Ø±Ø§Ø¬Ø¹ ÙˆÙÙ‚ Ø£Ø³Ù„ÙˆØ¨ APA Ø§Ù„Ø¹Ø§Ù„Ù…ÙŠ.</div>
        </div>""", unsafe_allow_html=True)
    with c3:
        st.markdown("""
        <div class="service-card">
            <span class="srv-icon">ğŸ”</span>
            <div class="srv-title">ØªØ­Ù„ÙŠÙ„ Ø§Ù„Ù…Ø±Ø§Ø¬Ø¹</div>
            <div class="srv-desc">ØªÙ„Ø®ÙŠØµ Ø§Ù„ÙƒØªØ¨ ÙˆØ§Ù„Ù…Ù‚Ø§Ù„Ø§Øª Ø§Ù„Ø·ÙˆÙŠÙ„Ø© ÙˆØ§Ø³ØªØ®Ø±Ø§Ø¬ Ø§Ù„Ø²Ø¨Ø¯Ø©.</div>
        </div>""", unsafe_allow_html=True)
    with c4:
        st.markdown("""
        <div class="service-card">
            <span class="srv-icon">âœ’ï¸</span>
            <div class="srv-title">Ø§Ù„ØªØ¯Ù‚ÙŠÙ‚ Ø§Ù„Ù„ØºÙˆÙŠ</div>
            <div class="srv-desc">ØªØµØ­ÙŠØ­ Ø§Ù„Ø£Ø®Ø·Ø§Ø¡ ÙˆØªØ­Ø³ÙŠÙ† Ø§Ù„Ø£Ø³Ù„ÙˆØ¨ Ø§Ù„Ø£ÙƒØ§Ø¯ÙŠÙ…ÙŠ Ù„Ù„Ù†Øµ.</div>
        </div>""", unsafe_allow_html=True)

    # --- CTA (Start Now) ---
    st.markdown("<br>", unsafe_allow_html=True)
    c_btn1, c_btn2, c_btn3 = st.columns([1, 2, 1])
    with c_btn2:
        if st.button("ğŸš€ Ø§Ø¨Ø¯Ø£ Ø§Ù„Ø¢Ù† Ù…Ø¬Ø§Ù†Ø§Ù‹", use_container_width=True):
            go_to_auth()

    # --- Contact Form ---
    st.markdown("<br><br>", unsafe_allow_html=True)
    with st.container():
        st.markdown("""
        <div class="contact-section">
            <h3 style="text-align:center; color:#0d47a1;">ğŸ“¬ ØªÙˆØ§ØµÙ„ Ù…Ø¹Ù†Ø§</h3>
            <p style="text-align:center; color:#666;">Ù„Ø¯ÙŠÙƒ Ø§Ø³ØªÙØ³Ø§Ø±ØŸ Ù†Ø­Ù† Ù‡Ù†Ø§ Ù„Ù„Ù…Ø³Ø§Ø¹Ø¯Ø©</p>
        </div>
        """, unsafe_allow_html=True)
        
        c_form1, c_form2, c_form3 = st.columns([1, 2, 1])
        with c_form2:
            with st.form("contact_us"):
                name = st.text_input("Ø§Ù„Ø§Ø³Ù… Ø§Ù„ÙƒØ§Ù…Ù„")
                msg = st.text_area("Ø±Ø³Ø§Ù„ØªÙƒ")
                if st.form_submit_button("Ø¥Ø±Ø³Ø§Ù„ Ø§Ù„Ø±Ø³Ø§Ù„Ø©"):
                    if name and msg: st.success("Ø´ÙƒØ±Ø§Ù‹ Ù„Ùƒ! ØªÙ… Ø§Ø³ØªÙ„Ø§Ù… Ø±Ø³Ø§Ù„ØªÙƒ.")
                    else: st.error("Ø§Ù„Ø±Ø¬Ø§Ø¡ Ù…Ù„Ø¡ Ø¬Ù…ÙŠØ¹ Ø§Ù„Ø­Ù‚ÙˆÙ„")

    st.markdown("<br><br><div style='text-align:center; color:#aaa; font-size:0.9rem;'>Ø¬Ù…ÙŠØ¹ Ø§Ù„Ø­Ù‚ÙˆÙ‚ Ù…Ø­ÙÙˆØ¸Ø© Â© 2025 Ø§Ù„Ù…Ø´Ø±Ù Ø§Ù„Ø§ÙØªØ±Ø§Ø¶ÙŠ</div>", unsafe_allow_html=True)
    st.stop()

# ==========================================
# ğŸ” Login Flow
# ==========================================
if not st.session_state.logged_in:
    if st.button("ğŸ”™ Ø§Ù„Ø¹ÙˆØ¯Ø© Ù„Ù„Ø±Ø¦ÙŠØ³ÙŠØ©"):
        st.session_state.page_state = "landing"
        st.rerun()

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("""<div style='background:white;padding:40px;border-radius:20px;text-align:center;box-shadow:0 10px 30px rgba(0,0,0,0.05); border-top: 5px solid #1565c0;'><h2 style='color:#1565c0;'>ØªØ³Ø¬ÙŠÙ„ Ø§Ù„Ø¯Ø®ÙˆÙ„</h2></div><br>""", unsafe_allow_html=True)
        t1, t2 = st.tabs(["Ø¯Ø®ÙˆÙ„", "Ø­Ø³Ø§Ø¨ Ø¬Ø¯ÙŠØ¯"])
        with t1:
            with st.form("log"):
                e = st.text_input("Ø§Ù„Ø¨Ø±ÙŠØ¯ Ø§Ù„Ø¥Ù„ÙƒØªØ±ÙˆÙ†ÙŠ"); p = st.text_input("ÙƒÙ„Ù…Ø© Ø§Ù„Ù…Ø±ÙˆØ±", type="password")
                if st.form_submit_button("Ø¯Ø®ÙˆÙ„"):
                    v, u = check_login(e, p)
                    if v: st.session_state.logged_in=True; st.session_state.user_info=u; st.rerun()
                    else: st.error("Ø¨ÙŠØ§Ù†Ø§Øª Ø®Ø§Ø·Ø¦Ø©")
        with t2:
            with st.form("reg"):
                n = st.text_input("Ø§Ù„Ø§Ø³Ù…"); e = st.text_input("Ø§Ù„Ø¨Ø±ÙŠØ¯"); p = st.text_input("ÙƒÙ„Ù…Ø© Ø§Ù„Ù…Ø±ÙˆØ±", type="password")
                if st.form_submit_button("Ø¥Ù†Ø´Ø§Ø¡ Ø­Ø³Ø§Ø¨"):
                    ok, m = register_user(e, p, n)
                    if ok: st.success("ØªÙ… Ø¥Ù†Ø´Ø§Ø¡ Ø§Ù„Ø­Ø³Ø§Ø¨!"); st.info("Ø³Ø¬Ù„ Ø¯Ø®ÙˆÙ„Ùƒ Ø§Ù„Ø¢Ù†.")
                    else: st.error("Ø®Ø·Ø£")
    st.stop()

# ==========================================
# ğŸ’° Paywall & Config
# ==========================================
curr_email = str(st.session_state.user_info.get('username')).lower()
try: curr_status = str(st.session_state.user_info.get('status')).lower()
except: curr_status = "expired"
is_admin = (curr_email == ADMIN_EMAIL.strip().lower())
is_active = is_admin or curr_status == 'active'

try: genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
except: st.stop()

@st.cache_resource
def get_model():
    m = [x.name for x in genai.list_models() if 'generateContent' in x.supported_generation_methods]
    return next((x for x in m if 'flash' in x), m[0])

# --- Sidebar ---
with st.sidebar:
    # --- IMPORTANT: Custom Close Button for Mobile ---
    if st.session_state.mobile_menu_open:
        if st.button("âœ– Ø¥ØºÙ„Ø§Ù‚ Ø§Ù„Ù‚Ø§Ø¦Ù…Ø©", key="close_menu_btn"):
            close_menu()
            st.rerun()
    
    status_color = "#2ecc71" if is_active else "#ef5350"
    status_text = "Ù†Ø´Ø·" if is_active else "ØºÙŠØ± Ù…ÙØ¹Ù„"
    st.markdown(f"<div style='background:{status_color};padding:10px;border-radius:8px;color:white;text-align:center;margin-bottom:20px;'><b>{st.session_state.user_info.get('name')}</b><br><small>{status_text}</small></div>", unsafe_allow_html=True)
    
    if is_admin:
        if st.checkbox("Ù„ÙˆØ­Ø© Ø§Ù„ØªØ­ÙƒÙ… (Admin)"): st.session_state.admin_mode = True
        else: st.session_state.admin_mode = False
    
    st.markdown("---")
    lang = st.selectbox("Ø§Ù„Ù„ØºØ© / Language", ["English", "FranÃ§ais", "Ø§Ù„Ø¹Ø±Ø¨ÙŠØ©"])
    T = UI_TEXT[lang]
    st.markdown(f"<style>.stApp {{ direction: {T['dir']}; text-align: {T['align']}; }} [data-testid='stPopover']::before {{ content: '{T['chat_title']}'; }} </style>", unsafe_allow_html=True)
    
    st.subheader(T["sidebar_settings"])
    field = st.selectbox(T["field_label"], T["fields"])
    level = st.radio(T["level_label"], T["levels"])
    
    t_names = list(T["tasks"].keys())
    task_disp = st.selectbox(T["task_label"], t_names)
    internal_task_key = T["tasks"][task_disp]
    
    if is_active:
        with st.expander(T["history_label"]):
            hist = get_history(curr_email)
            for i, h in enumerate(reversed(hist)):
                st.button(f"{h.get('date')} | {h.get('task').split(':')[0]}", key=f"h_{i}", on_click=set_archive, args=(h.get('content'), h.get('task')))
    
    if st.button("ØªØ³Ø¬ÙŠÙ„ Ø®Ø±ÙˆØ¬"): logout()

# --- Admin ---
if st.session_state.get("admin_mode", False) and is_admin:
    st.title("Ù„ÙˆØ­Ø© Ø§Ù„ØªØ­ÙƒÙ…"); sh = get_sheet("users"); st.dataframe(pd.DataFrame(sh.get_all_records()))
    with st.form("adm"):
        t = st.text_input("Email"); a = st.selectbox("Status", ["active", "expired"])
        if st.form_submit_button("ØªØ­Ø¯ÙŠØ«"):
            try: c = sh.find(t); sh.update_cell(c.row, 4, a); st.success("ØªÙ… Ø§Ù„ØªØ­Ø¯ÙŠØ«")
            except: st.error("Ø®Ø·Ø£")
    st.stop()

# ==========================================
# ğŸ’³ Payment Page
# ==========================================
if st.session_state.show_payment_page and not is_active:
    if st.button(T['cancel_btn']):
        st.session_state.show_payment_page = False
        st.rerun()

    st.markdown(f"<h1 style='text-align:center; color:#1565c0;'>{T['pay_title']}</h1>", unsafe_allow_html=True)
    st.markdown(f"""<div style="background:white;padding:20px;border-radius:15px;border-left:5px solid #ffca28;margin-bottom:30px;"><h3 style="margin:0;color:#0d47a1;">ğŸš€ {T['pay_pitch_title']}</h3><p style="margin-top:10px;color:#555;">{T['pay_pitch_body']}</p></div>""", unsafe_allow_html=True)
    
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown(f"<div class='plan-card'><h4>{T['plans']['1']}</h4><div class='price-tag'>999 DZD</div><div class='plan-desc'>{T['plan_desc']['1']}</div></div>", unsafe_allow_html=True)
        if st.button(f"{T['select_btn']} 1", key="p1"): st.session_state.selected_plan = "Monthly"
    with c2:
        st.markdown(f"<div class='plan-card' style='border:2px solid #ffca28;'><h4>{T['plans']['6']}</h4><div class='price-tag'>5000 DZD</div><div class='plan-desc'>{T['plan_desc']['6']}</div></div>", unsafe_allow_html=True)
        if st.button(f"{T['select_btn']} 6", key="p2"): st.session_state.selected_plan = "6 Months"
    with c3:
        st.markdown(f"<div class='plan-card'><h4>{T['plans']['12']}</h4><div class='price-tag'>10,000 DZD</div><div class='plan-desc'>{T['plan_desc']['12']}</div></div>", unsafe_allow_html=True)
        if st.button(f"{T['select_btn']} 12", key="p3"): st.session_state.selected_plan = "Yearly"

    if st.session_state.selected_plan:
        st.markdown("---")
        st.info(f"âœ… {st.session_state.selected_plan}")
        with st.form("confirm_pay"):
            st.write(f"### ğŸ’³ BaridiMob")
            st.markdown("""<h2 style='color:#0d47a1; background:#e3f2fd; padding:10px; border-radius:10px; text-align:center;'>00799999002283727175</h2>""", unsafe_allow_html=True)
            ref = st.text_input("Transaction Reference / Ø±Ù‚Ù… Ø§Ù„ÙˆØµÙ„")
            if st.form_submit_button("âœ… ØªØ£ÙƒÙŠØ¯ Ø§Ù„Ø¯ÙØ¹"):
                if ref:
                    submit_payment(curr_email, ref, st.session_state.selected_plan)
                    rain_graduation_caps()
                    st.success(T['pay_success'])
                    time.sleep(5)
                    st.session_state.show_payment_page = False
                    st.rerun()
                else: st.error(T['pay_error'])
    st.stop()

# --- Workspace ---
col_main, _ = st.columns([1, 0.01])
model = genai.GenerativeModel(get_model())
student_name = st.session_state.user_info.get('name')
base_prompt = f"Role: Academic Supervisor. Lang: {lang}. Field: {field}. Level: {level}. User: {student_name}. Persona: Helpful Mentor."

with col_main:
    # --- Global Header (Visible only when logged in) ---
    st.markdown("""
    <div class="global-header">
        <h1 class="main-title">Virtual Supervisor v2</h1>
        <div class="fixed-slogan">Research Smarter, Not Harder</div>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown(f"<div class='integrity-box'>âš ï¸ {T['warn_msg']}</div>", unsafe_allow_html=True)

    if st.session_state.res_restored:
        st.markdown(f"<div class='result-card' style='border-left:5px solid #ffca28'><h5>ğŸ“œ {st.session_state.restored_task}</h5><hr>{st.session_state.res_restored}</div>", unsafe_allow_html=True)
        st.button("Close Archive", on_click=close_archive)

    is_free_task = (internal_task_key == "structure")
    is_fully_free = (internal_task_key == "discuss_topic")
    
    if not is_active and not is_free_task and not is_fully_free:
        st.markdown(f"""<div class="result-card" style="text-align:center;border:2px solid #ef5350;"><h2 style="color:#ef5350;">ğŸš« {T['pay_msg']}</h2></div>""", unsafe_allow_html=True)
        if st.button(f"ğŸš€ {T['upgrade_btn']}", on_click=go_to_payment): pass
    else:
        st.header(f"ğŸ“ {task_disp}")
        with st.form("work_form"):
            u_inp = ""
            u_file = None
            
            # --- ğŸ”¥ ÙˆØ§Ø¬Ù‡Ø© Ø®Ø§ØµØ© Ù„ØªÙ†Ø³ÙŠÙ‚ Ø§Ù„Ù…Ø±Ø§Ø¬Ø¹ ---
            if internal_task_key == "formatting":
                u_inp = st.text_area(T["ref_ph"], height=200)
                # Ù‚Ø§Ø¦Ù…Ø© Ù…Ù†Ø³Ø¯Ù„Ø© Ù„Ø£Ù†Ø¸Ù…Ø© Ø§Ù„ØªÙˆØ«ÙŠÙ‚
                style = st.selectbox(T["format_label"], T["citation_styles"])
            
            elif internal_task_key == "analyze":
                u_file = st.file_uploader(T["file_ph"], type="pdf")
                u_inp = st.text_input("Question")
            else:
                u_inp = st.text_area(T["input_ph"], height=150)
            
            submitted = st.form_submit_button(T["exec_btn"], type="primary")
        
        if submitted:
            if u_inp or u_file:
                with st.spinner("Thinking..."):
                    final_p = ""
                    if internal_task_key == "discuss_topic":
                        final_p = f"Discuss feasibility: '{u_inp}'. GUARDRAIL: Discussion only."
                    elif internal_task_key == "structure":
                        final_p = f"Create detailed thesis structure. Write 1500 words. Topic: '{u_inp}'"
                    elif internal_task_key == "references":
                        final_p = f"Suggest 10 academic references (APA 7). Topic: '{u_inp}'"
                    
                    # --- ğŸ”¥ Ø§Ø³ØªØ®Ø¯Ø§Ù… Ø§Ù„Ù…ØªØºÙŠØ± style ---
                    elif internal_task_key == "formatting":
                        final_p = f"Reformat and organize this list of references strictly according to {style} style rules. Fix punctuation, italics, and ordering. Input:\n{u_inp}"
                    
                    elif internal_task_key == "proofread":
                        final_p = f"Academic proofreading. Text: '{u_inp}'"
                    elif internal_task_key == "analyze" and u_file:
                        pdf = PyPDF2.PdfReader(u_file); txt = "".join([p.extract_text() for p in pdf.pages[:10]])
                        final_p = f"Analyze content. Query: {u_inp}\nContext: {txt[:5000]}"
                    
                    try:
                        res = model.generate_content(base_prompt + "\n" + final_p)
                        
                        if not is_active and internal_task_key == "structure":
                            preview = res.text[:1500] 
                            st.markdown(f"""
                            <div class="result-card">
                                {preview}...
                                <div class="blur-content">
                                    <br><br><br><br>
                                    <div class="paywall-overlay">
                                        <div class="pay-btn-overlay">{T['upgrade_btn']}</div>
                                    </div>
                                </div>
                            </div>
                            """, unsafe_allow_html=True)
                            st.button(f"ğŸ”“ {T['upgrade_btn']}", key="blur_upg", on_click=go_to_payment)
                        else:
                            st.session_state.last_res = res.text
                            st.session_state.last_task = f"{task_disp}: {u_inp[:20]}"
                            st.session_state.current_context = f"TASK: {task_disp}\nINPUT: {u_inp}\nRESULT: {res.text}"
                            st.rerun()
                    except Exception as e: st.error(str(e))

    if st.session_state.last_res:
        st.markdown(f"<div class='result-card'>{st.session_state.last_res}</div>", unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1:
            if st.button(T["save_btn"]):
                save_research(curr_email, st.session_state.last_task, st.session_state.last_res)
                st.toast("Saved!", icon="âœ…")
        with c2:
            docx = create_word_docx(st.session_state.last_res, title=task_disp)
            st.download_button(T["dl_btn"], docx, "result.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# --- Chat ---
with st.popover("", use_container_width=False):
    st.markdown(f"### {T['chat_title']}")
    chat_c = st.container(height=400)
    for m in st.session_state.chat_history:
        with chat_c.chat_message(m["role"]): st.markdown(m["content"])
    if q := st.chat_input("..."):
        st.session_state.chat_history.append({"role":"user","content":q})
        chat_c.chat_message("user").write(q)
        if is_active:
            try:
                r = model.generate_content(f"{base_prompt}\nCTX:{st.session_state.current_context}\nQ:{q}")
                st.session_state.chat_history.append({"role":"assistant","content":r.text})
                chat_c.chat_message("assistant").write(r.text)
            except: pass
        else:
            msg = "ğŸ”’ Please upgrade to unlock chat support."
            st.session_state.chat_history.append({"role":"assistant","content":msg})
            chat_c.chat_message("assistant").write(msg)